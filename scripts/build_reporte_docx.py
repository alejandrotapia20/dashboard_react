#!/usr/bin/env python3
"""Genera Reporte_Proyecto02.docx a partir de Reporte_Proyecto02.md.

Estilos: '# ' -> Heading 1, subtítulos en negrita (Normal), tablas Markdown
-> tablas de Word con estilo de cuadrícula. Uso: python3 scripts/build_reporte_docx.py
"""
import re
from pathlib import Path
import docx
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent.parent
MD = ROOT / "Reporte_Proyecto02.md"
OUT = ROOT / "Reporte_Proyecto02.docx"

lines = MD.read_text(encoding="utf-8").splitlines()
doc = docx.Document()


def add_runs(paragraph, text):
    """Procesa **negritas** y *cursivas* dentro de un texto."""
    for part in re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            paragraph.add_run(part)


def emit_table(rows):
    # rows: lista de listas de celdas (la 2ª fila Markdown es el separador ---)
    header, *body = [r for i, r in enumerate(rows) if i != 1]
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for j, cell in enumerate(header):
        run = t.rows[0].cells[j].paragraphs[0].add_run(cell)
        run.bold = True
    for r in body:
        cells = t.add_row().cells
        for j, cell in enumerate(r):
            add_runs(cells[j].paragraphs[0], cell)
    doc.add_paragraph()


i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if not line or line.startswith("---"):
        i += 1
        continue

    # Tabla Markdown
    if line.startswith("|"):
        block = []
        while i < len(lines) and lines[i].lstrip().startswith("|"):
            cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
            block.append(cells)
            i += 1
        emit_table(block)
        continue

    # Encabezados
    if line.startswith("# "):
        doc.add_heading(line[2:].strip(), level=1)
    elif line.startswith("## "):
        doc.add_heading(line[3:].strip(), level=1)
    elif line.startswith("### "):
        p = doc.add_paragraph()
        p.add_run(line[4:].strip()).bold = True
    elif line.startswith("- "):
        add_runs(doc.add_paragraph(style="List Bullet"), line[2:].strip())
    elif re.match(r"^\d+\.\s", line):
        # Índice de contenido: lo dejamos como lista numerada simple
        add_runs(doc.add_paragraph(style="List Number"), re.sub(r"^\d+\.\s", "", line).strip())
    else:
        add_runs(doc.add_paragraph(), line)
    i += 1

# Tamaño base de fuente coherente
for p in doc.paragraphs:
    for run in p.runs:
        if run.font.size is None:
            run.font.size = Pt(11)

doc.save(OUT)
print(f"Generado: {OUT.name} ({len(doc.paragraphs)} párrafos)")
